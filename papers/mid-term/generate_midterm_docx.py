#!/usr/bin/env python3
"""Generate the mid-term assessment docx matching the university .doc template.
Section 二、硕士研究生论文工作阶段性总结 is populated from sections/*.md.
"""

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
SECTIONS_DIR = ROOT / "sections"
BIBLIO = ROOT / "references.bib"
CSL = ROOT / "csl" / "ieee.csl"


def set_run_font(run, font_name="宋体", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold


def format_cell_text(cell, font_name="宋体", size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        for run in paragraph.runs:
            set_run_font(run, font_name, size, bold)


def add_centered_heading(doc, text, font_name="黑体", size=18, bold=True, space_after=Pt(12)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p


def add_paragraph(doc, text, font_name="宋体", size=12, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  first_line_indent=Cm(0.74),
                  line_spacing=1.5,
                  space_after=Pt(6),
                  space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold)
    return p


def add_section_title(doc, text, font_name="黑体", size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold=True)
    return p


def add_blank_lines(doc, n=2):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5


def md_to_plain(md_path):
    """Convert a markdown file to plain text with numbered citations."""
    cmd = [
        "pandoc", str(md_path),
        "--bibliography", str(BIBLIO),
        "--citeproc",
        "--csl", str(CSL),
        "-t", "plain",
        "--wrap", "none",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, check=True)
    return result.stdout


def clean_plain_text(text):
    """Clean pandoc plain text for inclusion in the Word report."""
    # Remove figure blocks: [caption]\n\ncaption\n\n  (keeps only the caption line, then remove it)
    text = re.sub(r"\n\n\[([^\]]+)\]\n+\1\n+", "\n\n", text)
    # Remove standalone image caption lines that appear alone after block removal
    text = re.sub(r"\n+[^\n]*(?:图 |表 |Fig\. |Table )[^\n]{0,80}\n+", "\n", text)
    # Fix awkward ref removals: "如图 所示" -> "如图所示"
    text = re.sub(r"如图\s+所示", "如图所示", text)
    text = re.sub(r"图\s+展示了", "图展示了", text)
    text = re.sub(r"表\s+列出了", "表列出了", text)
    # Remove raw LaTeX refs if any remain
    text = re.sub(r"\\ref\{[^}]+\}", "", text)
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove leading/trailing whitespace
    return text.strip()


def split_into_paragraphs(text):
    """Split text into paragraphs, preserving list items."""
    paragraphs = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        paragraphs.append(stripped)
    return paragraphs


def create_info_table(doc):
    table = doc.add_table(rows=9, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    labels = [
        "姓    名",
        "学    号",
        "专    业",
        "导    师",
        "学    院",
        "研究方向",
        "论文题目",
        "开题报告日期",
        "论文中期考核    □ 第一次    □ 第二次",
    ]
    for i, label in enumerate(labels):
        row = table.rows[i]
        cell0 = row.cells[0]
        cell1 = row.cells[1]
        cell0.text = label
        cell0.width = Cm(3.5)
        cell1.width = Cm(10)
        format_cell_text(cell0, "宋体", 12)
        format_cell_text(cell1, "宋体", 12)
    return table


def create_plan_table(doc):
    categories = [
        ("课程学分", ["政治课程", "基础类课程", "素养类课程", "交叉、前沿类课程",
                     "专业基础类课程", "校企联合专业学位培养专项"]),
        ("非课程学分", ["科研育人", "学术活动", "课题组研讨活动", "文献阅读与评述",
                       "创新创业与社会实践（专业学位）", "专业实践（专业学位）", "开题报告"]),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    hdr[0].text = "个人培养计划"
    hdr[1].text = "计划学分"
    hdr[2].text = "完成学分"
    hdr[0].width = Cm(3)
    hdr[1].width = Cm(5.5)
    hdr[2].width = Cm(3.5)
    for cell in hdr:
        format_cell_text(cell, "宋体", 12, bold=True)

    for group, items in categories:
        start_idx = len(table.rows)
        for item in items:
            row = table.add_row().cells
            row[1].text = item
            row[0].width = Cm(3)
            row[1].width = Cm(5.5)
            row[2].width = Cm(3.5)
            for cell in row:
                format_cell_text(cell, "宋体", 12)
        if len(items) > 1:
            first_col_cells = [table.rows[start_idx + i].cells[0] for i in range(len(items))]
            first_col_cells[0].merge(first_col_cells[-1])
            first_col_cells[0].text = group
            format_cell_text(first_col_cells[0], "宋体", 12, bold=True)

    row = table.add_row().cells
    row[0].text = "总计"
    row[0].merge(row[1])
    row[0].width = Cm(8.5)
    row[2].width = Cm(3.5)
    format_cell_text(row[0], "宋体", 12, bold=True)
    format_cell_text(row[2], "宋体", 12)

    add_paragraph(doc, "备注：未完成的在“完成学分”栏填写“0”。", size=10.5,
                  first_line_indent=Cm(0), space_after=Pt(6))


def create_expert_table(doc):
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    hdr[0].text = "姓名"
    hdr[1].text = "是否硕导"
    hdr[2].text = "所在单位"
    widths = [Cm(3.5), Cm(3), Cm(6)]
    for cell, w in zip(hdr, widths):
        cell.width = w
        format_cell_text(cell, "宋体", 12, bold=True)
    for i in range(1, 5):
        for cell, w in zip(table.rows[i].cells, widths):
            cell.width = w
            format_cell_text(cell, "宋体", 12)
    return table


def add_section2_from_markdown(doc):
    """Populate section 2 from sections/report_*.md."""
    add_section_title(doc, "二、硕士研究生论文工作阶段性总结")
    add_paragraph(doc, "填写内容包括但不限于以下几点：", first_line_indent=Cm(0), size=12)
    add_paragraph(doc, "阐述论文工作的进展和取得的阶段性成果。", first_line_indent=Cm(0.74), size=12)
    add_paragraph(doc, "介绍论文发表的情况。", first_line_indent=Cm(0.74), size=12)
    add_paragraph(doc, "制定下一步研究计划及论文工作计划。", first_line_indent=Cm(0.74), size=12)

    # Read and clean each section
    intro = clean_plain_text(md_to_plain(SECTIONS_DIR / "report_introduction.md"))
    progress = clean_plain_text(md_to_plain(SECTIONS_DIR / "report_progress.md"))
    publication = clean_plain_text(md_to_plain(SECTIONS_DIR / "report_publication.md"))
    plan = clean_plain_text(md_to_plain(SECTIONS_DIR / "report_plan.md"))

    # Remap top-level headings to section-2 numbering
    intro = re.sub(r"^一、引言\s*", "", intro)
    intro = re.sub(r"^(\d+\.\d+)\s+", r"\1 ", intro)

    progress = re.sub(r"^二、工作进展与阶段性成果\s*", "", progress)
    publication = re.sub(r"^四、论文发表情况\s*", "", publication)
    publication = re.sub(r"\b4\.(\d+(?:\.\d+)*\s+)", r"3.\1", publication)
    plan = re.sub(r"^五、下一步研究计划与论文工作计划\s*", "", plan)
    plan = re.sub(r"\b5\.(\d+(?:\.\d+)*\s+)", r"4.\1", plan)

    # Combine
    parts = [
        ("1 研究背景与问题", intro),
        ("2 论文工作进展与阶段性成果", progress),
        ("3 论文发表情况", publication),
        ("4 下一步研究计划及论文工作计划", plan),
    ]

    for heading, body in parts:
        if not body.strip():
            continue
        add_paragraph(doc, heading, bold=True, first_line_indent=Cm(0.74), space_before=Pt(8))
        for para in split_into_paragraphs(body):
            if not para:
                continue
            # Detect markdown-like subsection headings (e.g., "1.1 研究背景")
            if re.match(r"^\d+(\.\d+)+\s+", para) and len(para) < 60:
                add_paragraph(doc, para, bold=True, first_line_indent=Cm(0.74), space_before=Pt(4))
            # Detect list items like "1. " or "- " or "（1）"
            elif re.match(r"^(\d+\.\s+|[-•]\s+|（\d+）)", para):
                add_paragraph(doc, para, first_line_indent=Cm(0.74))
            else:
                add_paragraph(doc, para, first_line_indent=Cm(0.74))

    add_blank_lines(doc, 4)


def main():
    doc = Document()
    sections = doc.sections[0]
    sections.page_height = Cm(29.7)
    sections.page_width = Cm(21.0)
    sections.top_margin = Cm(2.54)
    sections.bottom_margin = Cm(2.54)
    sections.left_margin = Cm(3.17)
    sections.right_margin = Cm(3.17)

    # Cover page
    add_blank_lines(doc, 1)
    add_centered_heading(doc, "西 南 交 通 大 学", font_name="宋体", size=22)
    add_centered_heading(doc, "硕士研究生论文中期考核报告", font_name="黑体", size=22, space_after=Pt(24))

    add_blank_lines(doc, 2)
    create_info_table(doc)
    add_blank_lines(doc, 3)

    add_paragraph(doc, "        年        月        日", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=Cm(0), size=12)
    add_blank_lines(doc, 2)
    add_paragraph(doc, "研究生院制表", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=Cm(0), size=12)

    doc.add_page_break()

    # 填表说明
    add_section_title(doc, "填表说明")
    notes = [
        "本表是硕士研究生论文工作中期考核用表。",
        "硕士研究生应在第四学期结束前撰写论文中期考核报告，汇报培养计划完成情况以及学术科研和论文工作的阶段性成果。",
        "“硕士研究生论文中期考核报告”经由导师签字同意后，将本表第一至第三项所在页面拍照制成PDF电子文档，上传研究生管理信息系统对应的模块中，纸质版交至学院。",
        "由导师（组）负责召集至少3名相关学科具有硕士研究生指导资格的专家组成的专家组对学生报告进行答辩评审。中期考核结果按“通过”或“不通过”记载。",
        "若首次论文中期考核不通过，硕士研究生应按照专家组意见进行整改后，可于下一学期再次进行中期考核。两次中期考核不通过者，由专家组作出“应予退学”处理建议报学校研究生院。",
        "每年秋季学期，院校两级督导组将对存档资料进行抽查。",
        "表格打印要求：A4纸双面打印。",
    ]
    for note in notes:
        add_paragraph(doc, note, size=12, first_line_indent=Cm(0))

    # 一、硕士研究生个人培养计划完成情况
    add_section_title(doc, "一、硕士研究生个人培养计划完成情况")
    add_paragraph(doc, "（请研究生根据《研究生成绩单》填写个人培养计划的比对结果）",
                  first_line_indent=Cm(0), size=12)
    create_plan_table(doc)

    # 二、硕士研究生论文工作阶段性总结（from sections/*.md）
    add_section2_from_markdown(doc)

    # 三、指导教师意见
    add_section_title(doc, "三、指导教师意见")
    add_blank_lines(doc, 8)
    add_paragraph(doc, "指导教师签字：", alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=Cm(0))
    add_paragraph(doc, "年    月    日", alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=Cm(0))

    # 四、专家组成员信息
    add_section_title(doc, "四、专家组成员信息")
    create_expert_table(doc)
    add_blank_lines(doc, 2)

    # 五、专家组考核结果
    add_section_title(doc, "五、专家组考核结果")
    add_blank_lines(doc, 6)
    add_paragraph(doc, "综上所述，该生本次论文中期考核结论：", first_line_indent=Cm(0))
    add_paragraph(doc, "□ 通过    □ 首次不通过（下学期再次考核）    □ 第二次考核不通过（建议退学）",
                  first_line_indent=Cm(0.74))
    add_blank_lines(doc, 4)
    add_paragraph(doc, "专家组成员签字：", alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=Cm(0))
    add_paragraph(doc, "年    月    日", alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line_indent=Cm(0))

    out_path = "硕士研究生论文中期考核报告_MACR.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
